"""
Minimal Markdown-to-Word renderer used to turn Claude's report text into a
downloadable .docx, per claude_skills_website_plan.md (python-docx, not docx-js).
Handles headings, bold/italic runs, bullet/numbered lists, simple pipe tables,
and horizontal rules — enough structure for the LLM-assisted report skills.
"""
import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

DARK_BLUE = RGBColor(0x2B, 0x57, 0x97)

_BOLD_ITALIC_RE = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)")


def _add_inline_runs(paragraph, text: str):
    pos = 0
    for match in _BOLD_ITALIC_RE.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        chunk = match.group(0)
        if chunk.startswith("***"):
            run = paragraph.add_run(chunk[3:-3])
            run.bold = True
            run.italic = True
        elif chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(chunk[1:-1])
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _flush_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = val.strip()
            if i == 0:
                for p in cells[j].paragraphs:
                    for r in p.runs:
                        r.bold = True
    doc.add_paragraph("")


def markdown_to_docx_bytes(title: str, subtitle: str, body_markdown: str) -> bytes:
    doc = Document()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = DARK_BLUE

    if subtitle:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        srun = sub_p.add_run(subtitle)
        srun.italic = True
        srun.font.size = Pt(13)
        srun.font.color.rgb = DARK_BLUE

    doc.add_paragraph("")

    lines = body_markdown.splitlines()
    table_buffer: list[list[str]] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        is_table_row = stripped.startswith("|") and stripped.endswith("|")
        if is_table_row:
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                table_buffer.append(cells)
            i += 1
            continue
        elif table_buffer:
            _flush_table(doc, table_buffer)
            table_buffer = []

        if not stripped:
            i += 1
            continue
        if stripped in ("---", "***", "___"):
            doc.add_paragraph("_" * 60)
            i += 1
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            h = doc.add_heading(level=min(level, 4))
            _add_inline_runs(h, text)
            for r in h.runs:
                r.font.color.rgb = DARK_BLUE
            i += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)", stripped)
        if bullet_match:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, bullet_match.group(1))
            i += 1
            continue

        numbered_match = re.match(r"^\d+[.)]\s+(.*)", stripped)
        if numbered_match:
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, numbered_match.group(1))
            i += 1
            continue

        p = doc.add_paragraph()
        _add_inline_runs(p, stripped)
        i += 1

    if table_buffer:
        _flush_table(doc, table_buffer)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
