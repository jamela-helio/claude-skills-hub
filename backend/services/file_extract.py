"""Best-effort plain-text extraction from uploaded planning/GIS/architectural files."""
import io
import zipfile


def extract_text_from_file(filename: str, raw: bytes) -> str:
    lower = filename.lower()
    try:
        if lower.endswith(".docx"):
            from docx import Document

            doc = Document(io.BytesIO(raw))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            return "\n".join(parts)

        if lower.endswith(".pdf"):
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw))
            return "\n".join(page.extract_text() or "" for page in reader.pages)

        if lower.endswith(".txt") or lower.endswith(".csv"):
            return raw.decode("utf-8", errors="ignore")

        return f"[{filename}: binary/image file — {len(raw)} bytes, not text-extracted]"
    except Exception as exc:  # noqa: BLE001
        return f"[{filename}: could not extract text — {exc}]"


def extract_zip_context(raw: bytes) -> str:
    """Unzip an uploaded archive and extract text from every member file."""
    sections = []
    with zipfile.ZipFile(io.BytesIO(raw)) as zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            member_bytes = zf.read(info.filename)
            text = extract_text_from_file(info.filename, member_bytes)
            sections.append(f"=== {info.filename} ===\n{text}")
    return "\n\n".join(sections)
